import docx
import io
import re
from verificadores import safe_eval

class DynamicDoc:

    def __init__(self):
        pass

    def load_document(self, document_path: str, **kwargs):
        self.doc = self.gera_docx(document_path, **kwargs)

    def handle_paragraphs(self,doc: docx.Document, vars: dict) -> None:
        paragraphs_to_delete = []
        lista = enumerate(doc.paragraphs)

        for i, p in lista:
            self.handle_paragraph(p, vars)
            if i in paragraphs_to_delete:
                p.clear()
                continue

            block_paragraphs = []
            if self.is_block(p):
                block_paragraphs.append(i)
                while not self.is_end_block(p):
                    i += 1
                    p = doc.paragraphs[i]
                    block_paragraphs.append(i)
                block_paragraphs.append(i)
                p.clear()
                paragraphs_to_delete.append(block_paragraphs[-1])

                self.handle_block([doc.paragraphs[j] for j in block_paragraphs], vars)
                               
    def is_block(self,p: any) -> bool:
        return re.search(r'>>if(.*?)<<', p.text) or re.search(r'>>for(.*?)<<', p.text) or re.search(r'>>elif(.*?)<<', p.text) or re.search(r'>>else<<', p.text)

    def is_end_block(self,p: any) -> bool:
        return re.search(r'>>endif<<', p.text) or re.search(r'>>endfor<<', p.text)

    def handle_block(self,list_of_paragraphs: list, vars: dict) -> list:
        #se for if 
        if re.search(r'>>if(.*?)<<', list_of_paragraphs[0].text):
            return self.handle_conditional_block(list_of_paragraphs, vars)

        #se for for
        elif re.search(r'>>for(.*?)<<', list_of_paragraphs[0].text):
            return self.handle_for(list_of_paragraphs, vars)

    def handle_conditional_block(self,list_of_paragraphs:list, vars:dict={}) -> list:
        # Verificar se a condição é verdadeira
        condition = str(re.search(r'>>if(.*?)<<', list_of_paragraphs[0].text).group(1)).strip()
        print("Condition: ", condition)
        if not safe_eval(condition):
            [p.clear() for p in list_of_paragraphs]
        list_of_paragraphs[0].clear()

    def handle_for(self,list_of_paragraphs:list, vars:dict={}) -> list:
        pass

    def handle_paragraph(self,p: any, vars: dict) -> None:
        for key in vars.keys():
            self.convert_code_to_text(key, vars[key], p)

    def convert_code_to_text(self,key: str, value:str, p) -> None:
        pattern_var = f">>{key}<<(\((N|I|N,I|I,N)\))?"
        pattern_if = f">>if ({key})?<<"
        pattern_for = f">>for ({key})?<<"
        pattern_elif = f">>elif ({key})?<<"
        
        pattern_var = f"{pattern_var}"
        other_patterns = f"{pattern_if}|{pattern_for}|{pattern_elif}"

        if re.search(pattern_var, p.text):
            match = re.search(pattern_var, p.text)
            
            if match:
                # Remover o "key" do texto
                text_before = p.text[:match.start()]
                text_after = p.text[match.end():]
                

                
                
                # Aplicar formatação dependendo do padrão encontrado
                run = p.clear()  # Limpa o parágrafo para refazer com formatação
                p.add_run(text_before)  # Adiciona o texto anterior ao padrão

                if "(N,I)" in match.group() or "(I,N)" in match.group():
                    formatted_run = p.add_run(str(value))
                    formatted_run.bold = True
                    formatted_run.italic = True
                elif "(N)" in match.group():
                    formatted_run = p.add_run(str(value))
                    formatted_run.bold = True
                elif "(I)" in match.group():
                    formatted_run = p.add_run(str(value))
                    formatted_run.italic = True
                else:
                    p.add_run(str(value))  # Sem formatação

                p.add_run(text_after)  # Adiciona o texto posterior ao padrão

        elif re.search(other_patterns, p.text):
            #modificando o key para o valor correspondente mas mantendo o >>if <<
            pattern = re.search(other_patterns, p.text).group(1)
            text_before = p.text[:re.search(pattern, p.text).start()]
            text_after = p.text[re.search(pattern, p.text).end():]
            run = p.clear()
            p.add_run(text_before)
            p.add_run(str(value))
            p.add_run(text_after)
        else:
            pass

    def gera_docx(self,document_path: str, **kwargs):
        doc = docx.Document(document_path)
        output = io.BytesIO()

        self.handle_paragraphs(doc, kwargs)
            
        doc.save(output)
        output.seek(0)
        
        return output.getvalue()

    def get_document(self):
        return self.doc

    def save_document(self, path: str):
        with open(path, 'wb') as f:
            f.write(self.doc)

    
if __name__ == "__main__":
    doc = DynamicDoc()
    doc.load_document("prep-4.docx", NOMES=False)
    doc.save_document("output.docx")

