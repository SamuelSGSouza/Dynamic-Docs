import os
import docx
from docx.text.paragraph import Paragraph
import io
import re
from verificadores import safe_eval

class DynamicDoc:

    def __init__(self, document_path: str = None, **kwargs):
        self.doc = self.gera_docx(document_path, **kwargs)
        self.html = ""

    def handle_paragraphs(self,doc: docx.Document, vars: dict) -> None:
        paragraphs_to_delete = []
        lista = enumerate(doc.paragraphs)

        for i, p in lista:
            self.handle_paragraph(p, vars)
            if i in paragraphs_to_delete:
                p.clear()
                continue
            block_paragraphs = []

            if self.is_block(p):
                block_paragraphs.append(i)
                while not self.is_end_block(p):
                    i += 1
                    p = doc.paragraphs[i]
                    block_paragraphs.append(i)
                block_paragraphs.append(i)
                p.clear()
                paragraphs_to_delete.append(block_paragraphs[-1])

                self.handle_block([doc.paragraphs[j] for j in block_paragraphs], vars)
                               
    def is_block(self,p: any) -> bool:
        '''
        Essa função verifica um paragrafo passado e verifica se ele tem um dos padrões
        de bloco, como "if", "elif", "else" ou "for".
        Caso tenha um dos padrões, verifica se esse padrão é a única coisa na linha. Se for,
        é um bloco de if, senão o if deve ser um if de linha a ser tratado em outra função
        '''

        padrao = re.search(r'>#if\s*(.*?)\s*#<', p.text) or re.search(r'>#for(.*?)#<', p.text) or re.search(r'>#elif(.*?)#<', p.text) or re.search(r'>#else#<', p.text) or None
        if padrao:
            if len(str(padrao.group(0)).strip()) == len(str(p.text).strip()):
                return True
        return False

    

    def is_end_block(self,p: any) -> bool:
        return re.search(r'>#endif#<', p.text) or re.search(r'>#endfor#<', p.text)

    def handle_block(self,list_of_paragraphs: list[Paragraph], vars: dict) -> list:
        #se for if 
        if re.search(r'>#if(.*?)#<', list_of_paragraphs[0].text):
            return self.handle_conditional_block(list_of_paragraphs, vars)

        #se for for
        elif re.search(r'>#for(.*?)#<', list_of_paragraphs[0].text):
            return self.handle_for(list_of_paragraphs, vars)

    def handle_conditional_block(self,list_of_paragraphs:list[Paragraph], vars:dict={}) -> list:
        # Verificar se a condição é verdadeira
        condition = str(re.search(r'>#if(.*?)#<', list_of_paragraphs[0].text).group(1)).strip()
        if not safe_eval(condition, self.kwargs):
            [p.clear() for p in list_of_paragraphs]
        list_of_paragraphs[0].clear()

    def handle_conditional_inline(self, text: str) -> str:
        '''
        Verifica se o texto tem condicionais de linha e, caso possua, resolve esses condicionais.
        '''
        pattern = re.search(r'>#if.*?#<', text)
        
        if pattern:
            inicio = ">#if"
            fim = ">#endif#<"
            if not re.search(r'>#endif\s*(.*?)\s*#<', text):
                text += fim
            
            # Encontra todos os blocos condicionais (if, elif, else)
            block_patterns = list(re.finditer(r'>#(if|elif|else)\s*(.*?)\s*#<', text))
            end_pattern = re.search(r'>#endif\s*(.*?)\s*#<', text)
            
            if not block_patterns or not end_pattern:
                return text
            
            # Extrai o texto entre o início e o fim do bloco condicional
            start_index = block_patterns[0].start()
            end_index = end_pattern.end()
            full_block = text[start_index:end_index]
            
            # Itera sobre os blocos condicionais
            kept_text = ""
            for i, match in enumerate(block_patterns):
                condition_type = match.group(1).strip()
                condition = match.group(2).strip() if condition_type != "else" else "True"
                
                # Avalia a condição
                if condition_type == "else" or safe_eval(condition, self.kwargs):
                    # Encontra o próximo bloco ou o fim do bloco condicional
                    next_block_start = block_patterns[i + 1].start() if i + 1 < len(block_patterns) else end_pattern.start()
                    kept_text = text[match.end():next_block_start]
                    break
            
            # Substitui o bloco condicional pelo texto mantido
            text = text[:start_index] + kept_text + text[end_index:]
            return text
        
        return text

    def handle_for(self,list_of_paragraphs:list[Paragraph], vars:dict={}) -> list:
        pass

    def handle_paragraph(self, p: Paragraph, vars: dict) -> None:
        # 1) Obter todo o texto do parágrafo
        original_text = p.text
        
        # 2) Definir o padrão para capturar qualquer chave do tipo >#chave#<
        #    Captura algo como >#solicitacao#<, >#nomeEntidade#<, etc.
        #    Grupo 1: nome da chave
        pattern = r">#\s*.*?\s*#<"

        if p.runs:
            runs_blocks = []
            for run in p.runs:
                run_text = self.handle_conditional_inline(run.text)

                # Encontrar todas as ocorrências no texto
                matches = re.findall(pattern, run_text)         
                for match in matches:
                    print(run_text)
                    # Nome da chave e formatação
                    key_name = match.replace(">#", "").replace("#<", "").strip()  # ex: 'solicitacao'
                    
                    # Valor real que queremos inserir
                    key_value = vars.get(key_name, f">#{key_name}#<")  
                    run_text = re.sub(r'>#\s*' + key_name + r'\s*#<', key_value, run_text)
                
                # Preservar quebras de linha dentro do run
                if '\n' in run_text:
                    lines = run_text.split('\n')
                    run.clear()
                    for i, line in enumerate(lines):
                        run.add_text(line)
                        if i < len(lines) - 1:
                            run.add_break()  # Adiciona uma quebra de linha
                else:
                    run.clear()
                    run.add_text(run_text)
                    

    def convert_code_to_text(self, key: str, value: str, p) -> None:
        pattern_var = f">#{key}#<(\((N|I|N,I|I,N)\))?"
        pattern_if = f">#if ({key})?#<"
        pattern_for = f">#for ({key})?#<"
        pattern_elif = f">#elif ({key})?#<"

        pattern_var = f"{pattern_var}"
        other_patterns = f"{pattern_if}|{pattern_for}|{pattern_elif}"

        while re.search(pattern_var, p.text):
            match = re.search(pattern_var, p.text)

            if match:
                # Remover o "key" do texto
                text_before = p.text[:match.start()]
                text_after = p.text[match.end():]

                # Limpar o parágrafo e adicionar o texto antes do padrão
                p.clear()
                p.add_run(text_before)  # Adiciona o texto anterior ao padrão

                # Aplicar formatação dependendo do padrão encontrado
                if "(N,I)" in match.group() or "(I,N)" in match.group():
                    formatted_run = p.add_run(str(value))
                    formatted_run.bold = True
                    formatted_run.italic = True
                elif "(N)" in match.group():
                    formatted_run = p.add_run(str(value))
                    formatted_run.bold = True
                elif "(I)" in match.group():
                    formatted_run = p.add_run(str(value))
                    formatted_run.italic = True
                else:
                    p.add_run(str(value))  # Sem formatação

                p.add_run(text_after)  # Adiciona o texto posterior ao padrão

        # Verificar outros padrões como if, for, etc.
        if re.search(other_patterns, p.text):
            pattern = re.search(other_patterns, p.text).group(1)
            text_before = p.text[:re.search(pattern, p.text).start()]
            text_after = p.text[re.search(pattern, p.text).end():]
            p.clear()
            p.add_run(text_before)
            p.add_run(str(value))
            p.add_run(text_after)
    
    def gera_docx(self,document_path: str, **kwargs):
        self.kwargs = kwargs
        output = io.BytesIO()
        if os.path.exists(document_path):
            doc = docx.Document(document_path)
            

            self.handle_paragraphs(doc, kwargs)
                
            doc.save(output)
            output.seek(0)
        
            return output.getvalue()
        else:
            raise Exception("File Not Found")


    def get_document(self):
        return self.doc

    def save_document(self, path: str):
        with open(path, 'wb') as f:
            f.write(self.doc)

